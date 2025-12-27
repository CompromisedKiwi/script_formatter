import os
import re
import tkinter
from tkinter import filedialog
from tkinter import messagebox
from docx import Document
from docx.shared import RGBColor


def colorize_after_colon(input_file, prefix):
    print(f"正在标红文件: {input_file}")
    RED = RGBColor(255, 0, 0) # 红色

    doc = Document(input_file)

    # 遍历所有段落
    for para in doc.paragraphs:
        if len(para.runs) == 0:
            continue
        full_text = para.text
        font_size = para.runs[0].font.size

        if full_text.startswith(('【', '△', '人物')):
            continue

        # 支持英文冒号和中文冒号
        parts = full_text.split('：', 1) # 1表示只分割一次
        if len(parts) == 1:
            parts = full_text.split(':', 1)
        if len(parts) == 1:
            continue

        before_text = parts[0] + '：'
        after_text = parts[1]

        # 清空原段落内容
        para.clear()

        # 添加冒号前的内容（保持默认颜色）
        before_run = para.add_run(before_text)
        before_run.font.size = font_size

        # 对冒号后的文本再做一次处理：将括号（含中文括号）及其内部内容单独提取，不标红加粗
        # 匹配中文括号（）和英文括号()
        pattern = re.compile(r'（[^）]*）')
        matches = pattern.finditer(after_text)
        last_pos = 0
        for m in matches:
            start, end = m.span()
            # 先添加括号前的内容（标红加粗）
            if start > last_pos:
                text_before = after_text[last_pos:start]
                after_run = para.add_run(text_before)
                after_run.font.color.rgb = RED
                after_run.font.bold = True
                after_run.font.size = font_size
            # 再添加括号本身及内部内容（不标红，不加粗）
            bracket_text = after_text[start:end]
            bracket_run = para.add_run(bracket_text)
            bracket_run.font.size = font_size
            last_pos = end
        # 处理剩余未匹配的文本（标红加粗）
        if last_pos < len(after_text):
            remaining_text = after_text[last_pos:]
            after_run = para.add_run(remaining_text)
            after_run.font.color.rgb = RED
            after_run.font.bold = True
            after_run.font.size = font_size

    # 保存结果
    output_file = os.path.join(os.path.dirname(input_file), prefix + '_' + os.path.basename(input_file))
    doc.save(output_file)
    print(f"    处理完成！已保存至: {output_file}")

if __name__ == '__main__':
    # 选取要处理的文件
    root = tkinter.Tk()
    root.withdraw()
    # 打开文件选择对话框，只允许选择文件，不允许选择文件夹
    file_path = filedialog.askopenfilename(
        title="请选择要标红对话的剧本文件",
        filetypes=[("文档", "*.docx"), ("文档", "*.doc")]
    )
    colorize_after_colon(file_path, '标红')
    messagebox.showinfo("格式化完成", f"格式化完成！已保存到: {os.path.dirname(file_path)}")

# 打包为可执行文件
# pyinstaller -D -w --version-file=version.txt --name "剧本格式化工具" .\dialogue_to_red.py